import os

from docx import Document
from docx.shared import Pt

from telegram import Update
from telegram.ext import ContextTypes

from ai import ask_ai
from prompts import GENERATOR_PROMPT

from database import (
    save_name,
    get_name,
    create_user,
    get_user,
    get_statistics,
    get_all_users,
    increment_consultations,
    increment_documents_created,
    increment_documents_checked,
    set_tariff,
    user_exists,
    get_tariff,
)

from keyboard import (
    main_keyboard,
    admin_keyboard
)

from law_search import (
    is_law_request,
    parse_law_query
)

from law_service import law_service
from admin import is_admin, broadcast_mode
from document_handlers import create_docx


# -------------------------
# Історія діалогів
# -------------------------

user_history = {}

# -------------------------
# Режим роботи користувача
# -------------------------

user_mode = {}
admin_mode = {}
document_context = {}



def create_docx(title: str, content: str) -> str:

    os.makedirs("generated", exist_ok=True)

    filename = "".join(
        c if c.isalnum() else "_"
        for c in title
    )[:40]

    path = f"generated/{filename}.docx"

    doc = Document()

    heading = doc.add_heading(title, level=1)

    for run in heading.runs:
        run.font.size = Pt(18)

    doc.add_paragraph(content)

    doc.add_paragraph(
        "\n--------------------------------"
    )

    footer = doc.add_paragraph(
        "Створено за допомогою LawAI"
    )

    for run in footer.runs:
        run.italic = True

    doc.save(path)

    return path


async def start(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE
):

    create_user(update.effective_user.id)

    await update.message.reply_text(
        "👋 Вітаю!\n\n"
        "Я LawAI — юридичний AI-помічник.\n\n"
        "Оберіть потрібний розділ нижче 👇",
        reply_markup=main_keyboard(
            is_admin=is_admin(
                update.effective_user.id
            )
        )
    )


async def message(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE
):

    user_id = update.effective_user.id

    create_user(user_id)

    text = update.message.text.strip()

    if user_id not in user_history:
        user_history[user_id] = []

    # ---------------------------------
    # Головне меню
    # ---------------------------------

    if text == "⚖️ Юридична консультація":

        user_mode[user_id] = "chat"

        await update.message.reply_text(
            "Опишіть вашу юридичну ситуацію."
        )

        return

    if text == "📄 Перевірити документ":

        user_mode[user_id] = "document"

        await update.message.reply_text(
            "Надішліть PDF або DOCX документ."
        )

        return

    if text == "📝 Створити документ":

        user_mode[user_id] = "generator"

        await update.message.reply_text(
            "Який документ потрібно створити?"
        )

        return

    if text == "📚 Закони України":

        await update.message.reply_text(
            "Пошук законів буде додано на наступному етапі."
        )

        return
        # ---------------------------------
    # Профіль
    # ---------------------------------

    if text == "👤 Мій профіль":

        user = get_user(user_id)

        if user is None:
            create_user(user_id)
            user = get_user(user_id)

        if not user["name"] or user["name"] == "Невідомо":

            await update.message.reply_text(
                "👤 Ви ще не завершили реєстрацію.\n\n"
                "Для реєстрації напишіть:\n\n"
                "👉 Мене звати (Ваше ім'я)."
            )

            return

        tariff_icon = "💎" if user["tariff"] == "PRO" else "⭐"

        await update.message.reply_text(
            "👤 <b>Ваш профіль</b>\n\n"

            f"🆔 <b>ID:</b> <code>{user_id}</code>\n"
            f"👤 <b>Ім'я:</b> {user['name']}\n"
            f"{tariff_icon} <b>Тариф:</b> {user['tariff']}\n"
            f"📅 <b>Дата реєстрації:</b> {user['register_date']}\n\n"

            "━━━━━━━━━━━━━━\n\n"

            f"💬 <b>Консультацій:</b> {user['consultations']}\n"
            f"📄 <b>Створено документів:</b> {user['documents_created']}\n"
            f"📑 <b>Перевірено документів:</b> {user['documents_checked']}\n\n"

            "━━━━━━━━━━━━━━\n\n"

            f"🌐 <b>Мова:</b> {user['language']}",

            parse_mode="HTML"
        )

        return
    
    # ---------------------------------
    # Видати PRO
    # ---------------------------------

    if text == "⭐ Видати PRO":

        if not is_admin(user_id):

            await update.message.reply_text(
                "⛔ У вас немає доступу."
            )

            return

        admin_mode[user_id] = "give_pro"

        await update.message.reply_text(
            "Введіть Telegram ID користувача:"
        )

        return

    # ---------------------------------
    # Забрати PRO
    # ---------------------------------

    if text == "❌ Забрати PRO":

        if not is_admin(user_id):

            await update.message.reply_text(
                "⛔ У вас немає доступу."
            )

            return

        admin_mode[user_id] = "remove_pro"

        await update.message.reply_text(
            "Введіть Telegram ID користувача:"
        )

        return
    # ---------------------------------
    # LawAI PRO
    # ---------------------------------
    if text == "⭐ LawAI PRO":

        user = get_user(user_id)

        if user is None:
            create_user(user_id)
            user = get_user(user_id)

        if user["tariff"] == "PRO":

            await update.message.reply_text(
                "💎 <b>LawAI PRO</b>\n\n"

                "Ваш тариф активний.\n\n"

                "Доступно:\n"
                "✅ Безлімітні юридичні консультації\n"
                "✅ Аналіз PDF\n"
                "✅ Аналіз DOCX\n"
                "✅ Аналіз фотографій\n"
                "✅ Генерація юридичних документів\n"
                "✅ Пріоритетна швидкість роботи\n\n"

                "Дякуємо, що користуєтесь LawAI ❤️",

                parse_mode="HTML"
            )

        else:

            await update.message.reply_text(
                "⭐ <b>LawAI FREE</b>\n\n"

                "Ваш поточний тариф: <b>FREE</b>\n\n"

                "Доступно:\n"
                "✅ Юридичні консультації - 100 повідомлень\n"
                "✅ Аналіз PDF - 3 PDF \n"
                "✅ Аналіз DOCX - 3 DOCX  \n"
                "✅ Аналіз фотографій - 3 фотографії \n"
                "✅ Генерація юридичних документів - 3 генерації\n\n"

                "━━━━━━━━━━━━━━\n\n"

                "💎 <b>LawAI PRO відкриє:</b>\n\n"

                "✔ Безлімітні консультації\n"
                "✔ Безлімітний аналіз документів\n"
                "✔ Максимальну швидкість\n"
                "✔ Доступ до нових функцій\n\n"

                "🚀 Найближчим часом стане доступна можливість оформити підписку PRO.",

                parse_mode="HTML"
            )

        return
    


    # ---------------------------------
    # Адмін-панель
    # ---------------------------------

    if text == "🔐 Адмін-панель":

        if not is_admin(user_id):

            await update.message.reply_text(
                "⛔ У вас немає доступу."
            )

            return
        

        stats = get_statistics()

        await update.message.reply_text(
            "👨‍💼 Панель адміністратора\n\n"
            f"👥 Користувачів: {stats['users']}\n"
            f"⭐ PRO: {stats['pro']}\n\n"
            f"💬 Консультацій: {stats['consultations']}\n"
            f"📄 Створено документів: {stats['documents_created']}\n"
            f"📑 Перевірено документів: {stats['documents_checked']}",
            reply_markup=admin_keyboard()
        )

        return
    
    if text == "📊 Статистика":

        if not is_admin(user_id):

            await update.message.reply_text(
                "⛔ У вас немає доступу."
            )

            return

        stats = get_statistics()

        await update.message.reply_text(
        "📊 <b>Статистика LawAI</b>\n\n"

        "━━━━━━━━━━━━━━\n\n"

        f"👥 <b>Користувачів:</b> {stats['users']}\n"
        f"⭐ <b>FREE:</b> {stats['users'] - stats['pro']}\n"
        f"💎 <b>PRO:</b> {stats['pro']}\n\n"

        "━━━━━━━━━━━━━━\n\n"

        f"💬 <b>Консультацій:</b> {stats['consultations']}\n"
        f"📄 <b>Створено документів:</b> {stats['documents_created']}\n"
        f"📑 <b>Перевірено документів:</b> {stats['documents_checked']}",

        parse_mode="HTML",
        reply_markup=admin_keyboard()
        )
        return

    if text == "📢 Розсилка":

        if not is_admin(user_id):

            await update.message.reply_text(
                "⛔ У вас немає доступу."
            )

            return

        broadcast_mode[user_id] = True

        await update.message.reply_text(
            "📢 Надішліть текст повідомлення для розсилки.\n\n"
            "Після цього повідомлення буде надіслано всім користувачам."
        )

        return
    
    # ---------------------------------
    # Список користувачів
    # ---------------------------------

    if text == "👥 Користувачі":

        if not is_admin(user_id):

            await update.message.reply_text(
                "⛔ У вас немає доступу."
            )

            return

        users = get_all_users()

        if not users:

            await update.message.reply_text(
                "Користувачів поки немає."
            )

            return

        message = "👥 Зареєстровані користувачі\n\n"

        for i, user in enumerate(users, start=1):

            message += (
                f"{i}. {user[1] or 'Невідомо'}\n"
                f"🆔 {user[0]}\n"
                f"⭐ {user[2]}\n"
                f"📅 {user[3] or '-'}\n\n"
            )

        message += (
            f"Всього користувачів: {len(users)}"
        )

        await update.message.reply_text(message)

        return
    
    # ---------------------------------
    # Реєстрація
    # ---------------------------------

    if text.lower().startswith("мене звати"):

        name = text[11:].strip().title()

        if name:

            save_name(user_id, name)

            await update.message.reply_text(
                f"Приємно познайомитися, {name}! 😊"
            )

        return


    # ---------------------------------
    # Моє ім'я
    # ---------------------------------

    if "як мене звати" in text.lower():

        name = get_name(user_id)

        if name:

            await update.message.reply_text(
                f"Тебе звати {name} 😊"
            )

        else:

            await update.message.reply_text(
                "Я ще не знаю твого імені."
            )

        return

        # ---------------------------------
    # Очікування ID для видачі PRO
    # ---------------------------------

    if admin_mode.get(user_id) == "give_pro":

        if not text.isdigit():

            await update.message.reply_text(
                "❌ Telegram ID повинен містити лише цифри."
            )

            return

        target_user = int(text)

        if not user_exists(target_user):

            await update.message.reply_text(
                "❌ Користувача з таким ID не знайдено."
            )

            admin_mode.pop(user_id)

            return

        if get_tariff(target_user) == "PRO":

            await update.message.reply_text(
                "ℹ️ Користувач уже має тариф PRO."
            )

            admin_mode.pop(user_id)

            return
        
        set_tariff(target_user, "PRO")

        await update.message.reply_text(
            f"✅ Користувачу {target_user} успішно видано тариф PRO."
        )

        try:

            await context.bot.send_message(
                chat_id=target_user,
                text=(
                    "🎉 Вітаємо!\n\n"
                    "Вам активовано тариф ⭐ LawAI PRO.\n\n"
                    "Дякуємо, що користуєтесь LawAI!"
                )
            )

        except Exception:
            pass

        admin_mode.pop(user_id)

        return
    
        # ---------------------------------
    # Очікування ID для забирання PRO
    # ---------------------------------

    if admin_mode.get(user_id) == "remove_pro":

        if not text.isdigit():

            await update.message.reply_text(
                "❌ Telegram ID повинен містити лише цифри."
            )

            return

        target_user = int(text)

        if not user_exists(target_user):

            await update.message.reply_text(
                "❌ Користувача з таким ID не знайдено."
            )

            admin_mode.pop(user_id)

            return

        if get_tariff(target_user) != "PRO":

            await update.message.reply_text(
                "ℹ️ Користувач не має тарифу PRO."
            )

            admin_mode.pop(user_id)

            return
        
        set_tariff(target_user, "FREE")

        await update.message.reply_text(
            f"✅ У користувача {target_user} успішно забрано тариф PRO."
        )

        try:

            await context.bot.send_message(
                chat_id=target_user,
                text=(
                    "ℹ️ Ваш тариф змінено.\n\n"
                    "Тепер у вас тариф FREE."
                )
            )

        except Exception:
            pass

        admin_mode.pop(user_id)

        return
    
    # ---------------------------------
    # Генерація документів
    # ---------------------------------

    if user_mode.get(user_id) == "generator":

        answer = ask_ai(
            [
                {
                    "role": "user",
                    "content": text
                }
            ],
            system_prompt=GENERATOR_PROMPT
        )

        docx_path = create_docx(
            "Юридичний документ",
            answer
        )

        with open(docx_path, "rb") as f:

            await update.message.reply_document(
                document=f,
                filename="LawAI_Document.docx",
                caption="📄 Документ готовий."
            )

        await update.message.reply_text(
            "✅ Документ успішно створено.\n\n"
            "📄 Завантажте файл DOCX вище."
        )

        os.remove(docx_path)

        increment_documents_created(user_id)

        user_mode[user_id] = "chat"

        return
    
        
        # ---------- Завершити роботу з документом ----------

    if text == "❌ Завершити роботу з документом":

        if user_id in document_context:

            document_context.pop(user_id)

            await update.message.reply_text(
                "✅ Роботу з документом завершено.\n\n"
                "Тепер можете поставити нове юридичне питання або завантажити інший документ."
            )

        else:

            await update.message.reply_text(
                "📄 Зараз немає відкритого документа."
            )

        return
        # ---------------------------------
            # ---------- Робота з останнім документом ----------

    if user_id in document_context:

        answer = ask_ai(
            [
                {
                    "role": "system",
                    "content":
                        "Ти досвідчений український юрист. "
                        "Відповідай лише на основі документа користувача. "
                        "Якщо користувач просить скласти відзив, заперечення, "
                        "клопотання, апеляційну чи касаційну скаргу або інший "
                        "процесуальний документ — сформуй його у повному вигляді."
                },
                {
                    "role": "user",
                    "content":
                        f"Документ:\n\n{document_context[user_id]}\n\n"
                        f"Запит:\n{text}"
                }
            ]
        )

        # Якщо користувач просить створити документ
        if any(word in text.lower() for word in [
            "напиши відзив",
            "склади відзив",
            "напиши клопотання",
            "склади клопотання",
            "напиши заяву",
            "склади заяву",
            "напиши позов",
            "склади позов",
            "апеляційну скаргу",
            "касаційну скаргу"
        ]):

            docx_path = create_docx(
                "Юридичний документ",
                answer
            )

            with open(docx_path, "rb") as f:

                await update.message.reply_document(
                    document=f,
                    filename="LawAI_Document.docx",
                    caption="📄 Документ готовий."
                )

            os.remove(docx_path)

        else:

            await update.message.reply_text(answer)

        return
    
    # Пошук законів
    # ---------------------------------
    if is_law_request(text):

        law = parse_law_query(text)

        result = law_service.get_article(
            law["article"],
            law["codex"]
        )

        if result is None:

            await update.message.reply_text(
                "❌ Статтю не знайдено."
            )

            return

        explanation = ask_ai(
            [
                {
                    "role": "user",
                    "content":
                        "Поясни простими словами:\n\n"
                        + result["text"]
                }
            ]
        )

        message = (
            f"📚 {result['codex']}\n\n"
            f"{result['text']}\n\n"
            f"💡 Пояснення:\n\n"
            f"{explanation}"
        )

        MAX = 4000

        for i in range(0, len(message), MAX):

            await update.message.reply_text(
                message[i:i + MAX]
            )

        return


    # ---------------------------------
    # Повернення в головне меню
    # ---------------------------------

    if text == "⬅️ Назад":

        await update.message.reply_text(
            "🏠 Головне меню",
            reply_markup=main_keyboard(
                is_admin=is_admin(user_id)
            )
        )

        return
    

        # ---------------------------------
    # Розсилка
    # ---------------------------------

    if user_id in broadcast_mode:

        users = get_all_users()

        success = 0
        failed = 0

        for user in users:

            try:

                await context.bot.send_message(
                chat_id=user[0],
                text=text
                )

                success += 1

            except Exception as e:

                print(user)
                print(e)

                failed += 1

        broadcast_mode.pop(user_id)

        await update.message.reply_text(
            "✅ Розсилку завершено.\n\n"
            f"👥 Надіслано: {success}\n"
            f"❌ Помилок: {failed}",
            reply_markup=admin_keyboard()
        )

        return
    
        # ---------------------------------
    # AI-консультація
    # ---------------------------------

    user_history[user_id].append(
        {
            "role": "user",
            "content": text
        }
    )

    answer = ask_ai(
        user_history[user_id][-20:]
    )

    user_history[user_id].append(
        {
            "role": "assistant",
            "content": answer
        }
    )

    increment_consultations(user_id)

    await update.message.reply_text(answer)
    