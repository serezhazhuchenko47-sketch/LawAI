import os

from telegram import Update
from telegram.ext import ContextTypes

from docx import Document
from docx.shared import Pt

from ai import ask_ai
from documents import extract_text
from prompts import DOCUMENT_PROMPT
from database import increment_documents_checked


def create_analysis_docx(filename: str, content: str) -> str:
    """
    Створює DOCX-звіт з юридичним аналізом.
    """

    os.makedirs("generated", exist_ok=True)

    safe_name = "".join(
        c if c.isalnum() or c in ("_", "-") else "_"
        for c in filename
    )

    path = os.path.join(
        "generated",
        f"Analysis_{safe_name}.docx"
    )

    doc = Document()

    title = doc.add_heading(
        "Юридичний аналіз документа",
        level=1
    )

    for run in title.runs:
        run.font.size = Pt(18)

    doc.add_paragraph(f"Документ: {filename}")
    doc.add_paragraph()
    doc.add_paragraph(content)
    doc.add_paragraph()

    footer = doc.add_paragraph(
        "Звіт сформовано LawAI"
    )

    for run in footer.runs:
        run.italic = True

    doc.save(path)

    return path


async def handle_document(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE
):

    try:

        document = update.message.document

        if document is None:
            return

        await update.message.reply_text(
            "📄 Документ отримано.\n\n"
            "🔍 Витягую текст..."
        )

        os.makedirs("temp", exist_ok=True)

        telegram_file = await context.bot.get_file(
            document.file_id
        )

        file_path = os.path.join(
            "temp",
            document.file_name
        )

        await telegram_file.download_to_drive(file_path)

        # Поки що не підтримуємо старий формат DOC
        if file_path.lower().endswith(".doc"):

            await update.message.reply_text(
                "⚠️ Формат DOC поки що не підтримується.\n\n"
                "Будь ласка, відкрийте документ у Microsoft Word "
                "або LibreOffice та збережіть його у форматі DOCX "
                "або PDF."
            )

            if os.path.exists(file_path):
                os.remove(file_path)

            return

        text = extract_text(file_path)

        if not text:

            await update.message.reply_text(
                "❌ Не вдалося отримати текст документа."
            )

            if os.path.exists(file_path):
                os.remove(file_path)

            return

        await update.message.reply_text(
            "⚖️ Аналізую документ..."
        )

        history = [
            {
                "role": "user",
                "content":
                    "Проаналізуй цей документ:\n\n" + text
            }
        ]

        answer = ask_ai(
            history,
            DOCUMENT_PROMPT
        )

        docx_path = create_analysis_docx(
            document.file_name,
            answer
        )

        with open(docx_path, "rb") as f:

            await update.message.reply_document(
                document=f,
                filename=os.path.basename(docx_path),
                caption="📄 Юридичний аналіз готовий"
            )

        await update.message.reply_text(
            "✅ Аналіз завершено.\n\n"
            "DOCX-звіт успішно сформовано."
        )

        increment_documents_checked(update.effective_user.id)

        if os.path.exists(docx_path):
            os.remove(docx_path)

        if os.path.exists(file_path):
            os.remove(file_path)

    except Exception as e:

        print("DOCUMENT ERROR:", e)

        try:
            await update.message.reply_text(
                f"❌ Помилка під час аналізу документа:\n{e}"
            )
        except Exception:
            pass